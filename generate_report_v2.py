from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path(__file__).parent; OUT=ROOT/'output'; OUT.mkdir(exist_ok=True)
FILE=OUT/'Fanimation_eProject_Report_v3.0.docx'
d=Document(); sec=d.sections[0]
for s in d.sections: s.top_margin=Inches(.8); s.bottom_margin=Inches(.75); s.left_margin=Inches(.8); s.right_margin=Inches(.8)
st=d.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10.5); st.paragraph_format.space_after=Pt(6); st.paragraph_format.line_spacing=1.15
for n,sz,col in [('Heading 1',16,'2E74B5'),('Heading 2',13,'2E74B5'),('Heading 3',11,'1F4D78')]:
 s=d.styles[n]; s.font.name='Calibri'; s.font.size=Pt(sz); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(col); s.paragraph_format.space_before=Pt(14); s.paragraph_format.space_after=Pt(6); s.paragraph_format.keep_with_next=True
d.styles['Caption'].font.size=Pt(9); d.styles['Caption'].font.italic=True
def h(t,l=1): d.add_heading(t,l)
def p(t='',style=None): return d.add_paragraph(t,style)
def b(t): return d.add_paragraph(t,style='List Bullet')
def shade(cell,color='E8EEF5'):
 e=OxmlElement('w:shd'); e.set(qn('w:fill'),color); cell._tc.get_or_add_tcPr().append(e)
def table(headers,rows):
 t=d.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
 for i,x in enumerate(headers):
  c=t.rows[0].cells[i]; c.text=x; shade(c)
  for r in c.paragraphs[0].runs: r.bold=True; r.font.size=Pt(9)
 for row in rows:
  cells=t.add_row().cells
  for i,x in enumerate(row):
   cells[i].text=str(x)
   for r in cells[i].paragraphs[0].runs: r.font.size=Pt(9)
 p('')
 p('Table '+str(table.n)+'. Summary of '+headers[0].lower()+' and related information.', 'Caption'); table.n+=1
table.n=1
def code(t):
 q=d.add_paragraph(); q.paragraph_format.left_indent=Inches(.2)
 for line in t.strip().splitlines():
  r=q.add_run(line+'\n'); r.font.name='Consolas'; r.font.size=Pt(8)
def diagram(title,source):
 h(title,3); p('Figure '+str(diagram.n)+'. '+title+'.', 'Caption'); diagram.n+=1; p('Mermaid diagram:', 'Caption'); code(source)
diagram.n=1
def placeholder(title,why):
 t=d.add_table(rows=1,cols=1); c=t.cell(0,0); shade(c,'F4F6F9'); c.text='[Screenshot placeholder]\n'+title
 c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
 p('Figure '+str(placeholder.n)+'. '+title+' - '+why, 'Caption'); placeholder.n+=1
placeholder.n=50
def page(): d.add_page_break()
def pgnum(par):
 par.alignment=WD_ALIGN_PARAGRAPH.RIGHT; par.add_run('Fanimation Report v2.0 | Page '); e=OxmlElement('w:fldSimple'); e.set(qn('w:instr'),'PAGE'); par._p.append(e)
sec.header.paragraphs[0].text='FANIMATION | ePROJECT REPORT v2.0'; sec.header.paragraphs[0].runs[0].font.size=Pt(8)
pgnum(sec.footer.paragraphs[0])

# Cover with group member spaces
for _ in range(5): p()
q=p(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=q.add_run('ePROJECT REPORT'); r.bold=True; r.font.size=Pt(14); r.font.color.rgb=RGBColor.from_string('146941')
q=p(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=q.add_run('FANIMATION'); r.bold=True; r.font.size=Pt(30); r.font.color.rgb=RGBColor.from_string('0B2545')
q=p(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.add_run('A Simple React Fan Catalogue Website').font.size=Pt(15)
p(); table(['Group member','Student ID'],[['Nwosu Kingsley Ikem','Student1639745'],['Jerry Ochuko Oghoghomeh','Student1727456'],['Tochukwu Oluwadarsimi Agusiobo','Student1726887'],['Centre','Aptech Gwarimpa'],['School','APTECH'],['Course code','ADSE/CPISM/MERN'],['Academic session','2026'],['Submission date','20 July 2026']])
p('This report is based on the Fanimation project source code that was provided. It explains only what is in the project.')
page()
h('Table of Contents'); p('Acknowledgements\n1. eProject Synopsis\n2. eProject Analysis\n3. eProject Design\n4. DFDs, Flowcharts and Process Diagrams\n5. Database Design / Structure\n6. Screenshot Placeholders\n7. Source Code with Comments\n8. User Guide\n9. Developer Guide\n10. Module Descriptions\n11. Testing\n12. Conclusion and Future Improvements')
p('This visible contents list is included so the report never opens with an empty TOC. All report headings use Word Heading styles for easy future editing.')
page()
h('Acknowledgements'); p('We thank our lecturers and classmates for their support during this project. We also thank everyone who gave us advice while we were learning React and web development. This report explains our work in a clear and simple way.')

h('1. eProject Synopsis')
h('1.1 Introduction',2); p('Fanimation is a one-page website for showing fan products and accessories. A user can look through products, filter them, open product details, add items to a temporary cart, and use a contact form.')
h('1.2 Purpose, Objectives and Scope',2)
for x in ['Show 15 fan and accessory products in a neat catalogue.','Let users filter by category, type, price, colour and brand.','Let users sort products, view details, save favourites and use a simple cart.','Show gallery, offers, about, FAQ and contact sections.']: b(x)
p('The project is front-end only. It has no login, database, payment, real checkout, product search bar, or message sending service.')
h('1.3 Technologies Used',2); table(['Tool','How it is used'],[['React','Builds the user interface with components and state.'],['Vite','Runs and builds the React project.'],['Bootstrap','Helps with layout, forms, grid and FAQ accordion.'],['Lucide React','Shows icons.'],['CSS','Styles every component and makes the page responsive.']])
h('1.4 Expected Outcome',2); p('The finished website gives a simple shopping-style experience where visitors can discover Fanimation products without leaving the page.')

h('2. eProject Analysis')
h('2.1 Existing and Proposed System',2); p('There is no older system inside the uploaded files. The proposed system is the website that was built: a front-end product catalogue. It solves the problem of showing many fan types in one place with filters instead of making users check products one by one.')
h('2.2 Requirements Found in the Code',2); table(['Requirement','Status'],[['Navigation between page sections','Implemented with smooth in-page links.'],['Product catalogue','Implemented with 15 local product records.'],['Filters and sorting','Implemented.'],['Product details modal','Implemented.'],['Favourites and cart','Implemented but only kept in React memory.'],['Gallery, offers, FAQ, about and contact','Implemented.'],['Contact validation','Implemented; the form does not send data.'],['Real checkout, login, database and search','Not implemented.']])
h('2.3 Non-functional Requirements',2); p('The site is responsive, uses simple labelled form fields, has alt text on images, and uses Bootstrap plus custom CSS. It loads product data locally, so there is no API delay in the current version.')
h('2.4 Assumptions, Constraints and Feasibility',2); p('Prices are stored as text like 95,000 and changed into numbers when needed. The cart and favourites are lost when the page refreshes. The project is easy to run as a demo because it does not need a server. For a real shop, it would need a backend, database and payment system.')

h('3. eProject Design')
h('3.1 Architecture',2); p('The browser opens index.html. main.jsx starts React and renders App.jsx. App puts all the page sections together and keeps shared cart and favourite data. ProductSection reads the product data and handles filters.')
diagram('Figure 1. Simple system architecture','''flowchart TD
Browser --> main[src/main.jsx] --> App
App --> Navbar & Hero & ProductSection & CartDrawer & Gallery & Offers & About & FAQ & Contact & Footer
ProductSection --> ProductFilter & ProductGrid
ProductGrid --> ProductCard --> ProductModal
Data[(products.js)] --> ProductSection''')
h('3.2 Folder Structure',2); code('''src/
  components/     page parts such as Navbar, Gallery and Contact
  data/products.js  15 product objects and price helper
  assets/images/    logo, hero image and product images
  styles/global.css global page rules
  App.jsx           main page and shared state
  main.jsx          React start file''')
h('3.3 Navigation, State and Data Flow',2); p('Navbar and footer links scroll to sections. Category links send a small browser event to ProductSection so it can select the right category. App holds favouriteIds, cartItems and isCartOpen. ProductSection holds filter values. Other components keep their own small state, for example the gallery image and contact form.')
h('3.4 Filtering and Modal Design',2); p('The filter checks category, type, colour, brand and maximum price. It then sorts by rating, newest product ID, or price. Clicking View Details opens a React modal. The modal closes with the close button, the background, or the Escape key.')
diagram('Figure 2. Product filtering flow','''flowchart TD
User --> Filter[Choose filter or sort]
Filter --> State[Update filter state]
State --> Check[Check local product list]
Check --> Sort[Sort matching products]
Sort --> Grid[Show cards or no-products message]''')
h('3.5 Responsive Design and Styling',2); p('Bootstrap grid classes control most layouts. Product cards use three columns on large screens and two on medium screens. The gallery moves from four to three to two columns as the screen becomes smaller. The mobile navigation uses an overlay panel.')

h('4. DFDs, Flowcharts and Process Diagrams')
diagram('4.1 Context DFD','''flowchart LR
Visitor -->|clicks, filters, form input| Site[Fanimation React website]
Site -->|products and feedback| Visitor
Products[(Local products.js)] --> Site
Site <--> Storage[(Browser storage for visitor count)]''')
diagram('4.2 Level 1 DFD','''flowchart TD
Visitor --> Navigation --> Sections
Visitor --> Catalogue --> Products[(products.js)]
Catalogue --> ProductCards
Visitor --> CartFavorites --> AppState[(React state)]
Visitor --> Contact --> Validation --> Message[On-screen success message]''')
diagram('4.3 Product details flowchart','''flowchart TD
Click[Click View Details] --> Modal[Open product modal]
Modal --> Choice{Choose action}
Choice -->|Add to Cart| Cart[Update cart and open drawer]
Choice -->|Heart| Favourite[Update favourite list]
Choice -->|Close / Escape / backdrop| End[Close modal]''')
diagram('4.4 Contact and full website process','''flowchart TD
Start --> Browse[Browse or filter products]
Browse --> Details[View details or add to cart]
Start --> Other[Gallery, offers, FAQ or about]
Start --> Contact[Fill contact form]
Contact --> Validate{Valid?}
Validate -->|No| Errors[Show errors]
Validate -->|Yes| Confirm[Show local confirmation]
Confirm --> Note[No message is sent in this version]''')
diagram('4.5 Application startup flowchart','''flowchart TD
Open[Open website] --> HTML[index.html loads]
HTML --> Main[main.jsx loads React, Bootstrap and CSS]
Main --> App[App renders all sections]
App --> Ready[Website ready for the visitor]''')
diagram('4.6 Homepage navigation flowchart','''flowchart TD
Click[Click navbar or footer link] --> Check{Product category?}
Check -->|Yes| Filter[Send category event]
Check -->|No| Target[Find page section]
Filter --> Target
Target --> Scroll[Smooth scroll to section]''')
diagram('4.7 Product viewing and modal flow','''flowchart TD
Grid[Product grid] --> Card[Product card]
Card --> Details[View Details]
Details --> Modal[Show image, rating, price and description]
Modal --> Close[Close modal or take cart/favourite action]''')
diagram('4.8 Product search/filtering process','''flowchart LR
Note[Text search is not implemented] --> Filters[Category, type, price, colour and brand filters]
Filters --> Match[Match local product records]
Match --> Sort[Sort result]
Sort --> Display[Display cards]''')
diagram('4.9 Cart and contact process','''flowchart TD
Add[Add product] --> Cart[Update temporary cart state] --> Drawer[Show cart drawer]
Form[Submit contact form] --> Valid{Fields valid?}
Valid -->|No| Errors[Show errors]
Valid -->|Yes| Success[Show local success message]
Success --> NoSend[No backend submission in current project]''')

h('5. Database Design / Structure')
p('Current project: there is no database. Product data is inside products.js. The design below is only a future production idea.')
diagram('5.1 Proposed ER Diagram','''erDiagram
CATEGORY ||--o{ PRODUCT : has
PRODUCT ||--o{ REVIEW : receives
USER ||--o{ REVIEW : writes
USER ||--o{ WISHLIST : saves
PRODUCT ||--o{ WISHLIST : is_saved
USER ||--o{ ORDER : places
ORDER ||--|{ ORDER_ITEM : contains
PRODUCT ||--o{ ORDER_ITEM : ordered
USER ||--o{ CONTACT_MESSAGE : sends''')
table(['Table','Important fields','Simple rules'],[['Categories','category_id PK, name','Name should be unique.'],['Products','product_id PK, category_id FK, name, price, description','Price must not be negative.'],['Users','user_id PK, name, email','Email should be unique.'],['Reviews','review_id PK, product_id FK, user_id FK, rating','Rating should be 1 to 5.'],['Contact Messages','message_id PK, name, email, message','Save messages sent from the form.'],['Wishlist','user_id FK, product_id FK','One saved product per user.'],['Orders / Order Items','order_id PK; product_id FK; quantity; unit_price','Needed for real checkout.']])
p('The future design is normalized because products, users, reviews and orders are stored in separate tables. This avoids repeating the same data many times.')

h('6. Screenshot Placeholders')
p('Take these screenshots after running the website. They are placeholders because the source upload did not include final screenshots.')
for t,x in [('Homepage','First view of the complete website.'),('Hero Banner','Hero text and Explore Collection button.'),('Navigation','Desktop navbar and navigation links.'),('Product Section','Product list with count and sorting.'),('Product Filters','Category, type, price, colour and brand controls.'),('Product Cards','Product image, price, rating and actions.'),('Product Details Modal','Details popup for one product.'),('Cart Drawer','Cart items, quantity buttons and subtotal.'),('Gallery','Product image gallery.'),('Offers','Discounted products with old and new prices.'),('About','About section and statistics.'),('FAQ','Expanded FAQ answer.'),('Contact','Contact details and form.'),('Footer','Footer links and contact details.'),('Tablet View','Medium-size responsive layout.'),('Mobile View','Small-screen layout and mobile navigation.')]: placeholder(t,x)

h('7. Source Code with Comments')
p('The code already has useful short comments. These are the main parts a beginner should understand.')
h('7.1 Cart logic',2); code('''setCartItems((prev) => {
  const existing = prev.find((item) => item.id === normalizedProduct.id);
  if (existing) {
    return prev.map((item) => item.id === normalizedProduct.id
      ? { ...item, quantity: item.quantity + 1 } : item);
  }
  return [...prev, { ...normalizedProduct, quantity: 1 }];
});'''); p('This checks if the product is already in the cart. If it is, quantity goes up. If not, a new cart item is added.')
h('7.2 Filter logic',2); code('''if (filters.category && product.category !== filters.category) return false;
if (filters.color && product.color !== filters.color) return false;
if (filters.price && toNumber(product.price) > Number(filters.price)) return false;'''); p('A product is hidden when it does not match a selected filter. toNumber changes a formatted price into a number for comparison.')
h('7.3 Contact validation',2); code('''if (!values.name.trim()) nextErrors.name = "Please enter your name.";
else if (!/^\\S+@\\S+\\.\\S+$/.test(values.email))
  nextErrors.email = "Please enter a valid email address.";'''); p('The form checks for name, email and message before showing its local success message.')

h('8. User Guide')
table(['Task','How to do it'],[['Move around the site','Use the navigation bar, footer links or Explore Collection button.'],['Find a product','Choose a category or use the filters and Sort By menu.'],['See more information','Click View Details. Close with X, background click or Escape.'],['Save / add items','Use the heart for favourite and Add to Cart in the modal.'],['Change cart quantity','Open the cart icon and use minus or plus.'],['Contact the team','Enter name, valid email and message, then Send Message.'],['Important note','The cart, checkout and contact form are demo features only. Data is not saved or sent.']])

h('9. Developer Guide')
h('9.1 Start the Project',2); code('''npm install
npm run dev
npm run build
npm run lint'''); p('Use npm run dev while developing. Main dependencies are React, Vite, Bootstrap and Lucide React.')
h('9.2 Add a Product',2); p('In src/data/products.js, add a new product object with a unique id, name, category, type, colour, price, rating, reviews, image and description. If it has oldPrice, it will also show in Offers.')
h('9.3 Maintain and Extend',2); p('Keep styles in the correct component CSS file. Add a backend before adding real checkout or message sending. Add testing for filters, cart, modal and contact validation. Product search is not in the project yet and should be added as a new feature, not described as existing.')
h('9.4 Build Note',2); p('During source checking, npm run build failed with a Vite/Rolldown emitted-name error and npm run lint reported a React effect warning in Navbar.jsx. These should be fixed before deployment.')

h('10. Module Descriptions')
table(['Module','What it does'],[['main.jsx','Starts React, Bootstrap and global CSS.'],['App','Shows all sections and keeps shared cart/favourite state.'],['Navbar','Navigation, mobile menu, visitor count and cart/favourite indicators.'],['Ticker','Scrolling announcements and live clock.'],['Hero','Hero image and button to products.'],['Features','Four simple feature cards.'],['ProductSection','Owns filtering and sorting logic.'],['ProductFilter','Shows controlled filter controls.'],['ProductGrid','Maps matching products into cards.'],['ProductCard','Shows a short product summary and opens modal.'],['ProductModal','Shows full product details and actions.'],['CartDrawer','Shows cart items, quantities and subtotal.'],['Gallery','Shows eight local product images in a lightbox.'],['Offers','Shows products with oldPrice values.'],['About','Shows company text and statistics.'],['FAQ','Shows Bootstrap FAQ accordion.'],['Contact','Validates the local contact form.'],['Footer','Shows links, categories and contact details.'],['products.js','Stores 15 product records and toNumber helper.'],['Component CSS files','Style each module and responsive views.']])

h('11. Testing')
table(['Test','Expected result','Result from review'],[['Category filter','Only selected category appears.','Implemented in ProductSection filter logic.'],['Sort','Cards change order by rating, id or price.','Implemented.'],['Modal','Open and close correctly.','Implemented with Escape/backdrop support.'],['Cart','Add, increase, decrease and calculate subtotal.','Implemented in React state.'],['Contact','Show errors or local success message.','Implemented; no data sending.'],['Responsive layout','Page adapts to smaller screens.','CSS/Bootstrap rules found; full browser test still needed.'],['Build','Project should build.','Failed in checked environment; fix before deployment.'],['Lint','No errors.','One Navbar React-hooks error found.']])
p('There are no automated test files in the project. Manual browser testing on Chrome, Firefox and Edge is recommended after fixing build and lint problems.')

h('12. Conclusion and Future Improvements')
p('Fanimation meets its main goal as a simple front-end fan catalogue. It shows products clearly and gives useful interactions like filtering, modal details, favourites, cart, gallery, offers and contact validation. We learned how React components, props, state and CSS work together in one website.')
for x in ['Fix the build and lint issues.','Add a backend, database, login and real contact-message sending.','Save cart and favourites for users.','Add real checkout, payment and orders.','Add product search and automated tests.','Replace placeholders with group details and real screenshots before submission.']: b(x)

settings=d.settings.element; e=OxmlElement('w:updateFields'); e.set(qn('w:val'),'true'); settings.append(e)
d.core_properties.title='Fanimation eProject Report v2.0'; d.save(FILE); print(FILE)
