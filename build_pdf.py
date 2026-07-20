from pathlib import Path
from docx import Document
import textwrap

root=Path(__file__).parent
src=root/'output'/'Fanimation_eProject_Report.docx'
out=root/'output'/'Fanimation_eProject_Report.pdf'
doc=Document(src)

# Minimal dependency-free PDF writer. It is intentionally based on the generated
# report so the PDF remains a direct content equivalent when LibreOffice is absent.
pages=[]; lines=[]; y=760
def esc(s): return s.replace('\\','\\\\').replace('(','\\(').replace(')','\\)')
def flush():
    global lines,y
    if lines: pages.append(lines)
    lines=[]; y=760
def put(text, size=9, bold=False, gap=12):
    global y
    font='F2' if bold else 'F1'
    maxchars=78 if size<=10 else 64
    for line in textwrap.wrap(text, width=maxchars, replace_whitespace=True, drop_whitespace=True) or ['']:
        if y<55: flush()
        lines.append(f'BT /{font} {size} Tf 45 {y} Td ({esc(line)}) Tj ET')
        y-=gap
for para in doc.paragraphs:
    text=para.text.strip()
    if not text: continue
    style=para.style.name
    if style.startswith('Heading 1'): put(text,15,True,20); y-=5
    elif style.startswith('Heading 2'): put(text,12,True,16)
    elif style.startswith('Heading 3'): put(text,10,True,14)
    elif style=='Caption': put(text,8,False,10)
    else: put(text,9,False,12)
for t in doc.tables:
    # Tables are appended as readable records after prose; their content remains searchable.
    for row in t.rows:
        put(' | '.join(cell.text.replace('\n',' ').strip() for cell in row.cells),8,False,10)
flush()
objects=[]
def add(b): objects.append(b); return len(objects)
f1=add(b'<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>')
f2=add(b'<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica-Bold >>')
page_ids=[]
for page_no, content in enumerate(pages, 1):
    footer=f'BT /F1 8 Tf 450 28 Td (Fanimation eProject Report | Page {page_no}) Tj ET'
    stream=('\n'.join(content+[footer])+'\n').encode('latin-1','replace')
    cid=add(b'<< /Length '+str(len(stream)).encode()+b' >>\nstream\n'+stream+b'endstream')
    page_ids.append(add(b'<< /Type /Page /Parent PAGES /MediaBox [0 0 612 792] /Resources << /Font << /F1 '+str(f1).encode()+b' 0 R /F2 '+str(f2).encode()+b' 0 R >> >> /Contents '+str(cid).encode()+b' 0 R >>'))
kids=b' '.join(str(i).encode()+b' 0 R' for i in page_ids)
pages_id=add(b'<< /Type /Pages /Kids [ '+kids+b' ] /Count '+str(len(page_ids)).encode()+b' >>')
for i in page_ids:
    objects[i-1]=objects[i-1].replace(b'/Parent PAGES',b'/Parent '+str(pages_id).encode()+b' 0 R')
catalog=add(b'<< /Type /Catalog /Pages '+str(pages_id).encode()+b' 0 R >>')
data=b'%PDF-1.4\n%\xe2\xe3\xcf\xd3\n'; offsets=[0]
for i,obj in enumerate(objects,1):
    offsets.append(len(data)); data+=str(i).encode()+b' 0 obj\n'+obj+b'\nendobj\n'
xref=len(data); data+=b'xref\n0 '+str(len(objects)+1).encode()+b'\n0000000000 65535 f \n'
for off in offsets[1:]: data+=f'{off:010d} 00000 n \n'.encode()
data+=b'trailer\n<< /Size '+str(len(objects)+1).encode()+b' /Root '+str(catalog).encode()+b' 0 R >>\nstartxref\n'+str(xref).encode()+b'\n%%EOF'
out.write_bytes(data)
print(out)
